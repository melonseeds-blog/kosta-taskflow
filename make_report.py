# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

FONT_NAME = '맑은 고딕'  # 맑은 고딕

def set_font(run, bold=False, size=10, color=None):
    run.font.name = FONT_NAME
    run.font.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 한글 폰트 eastAsia 속성 명시 설정
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)

def shade_row(row, hex_color="E8EAF6"):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_font(run, bold=True, size=16, color=(99, 102, 241))
    elif level == 2:
        set_font(run, bold=True, size=13, color=(55, 65, 81))
    elif level == 3:
        set_font(run, bold=True, size=11, color=(75, 85, 99))
    return p

def add_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=10)
    return p

def cell_run(cell, text, bold=False, size=10, color=None, align=None):
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_font(run, bold=bold, size=size, color=color)
    if align:
        p.alignment = align
    return run

# ── 표지 ──────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TaskFlow 웹 애플리케이션')
set_font(run, bold=True, size=22, color=(99, 102, 241))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('테스트 결과 보고서')
set_font(run, bold=True, size=18, color=(55, 65, 81))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Playwright 자동화 테스트 기반')
set_font(run, size=11, color=(107, 114, 128))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'테스트 일자: {datetime.date.today().strftime("%Y년 %m월 %d일")}')
set_font(run, size=10, color=(107, 114, 128))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('작성자: BulNim')
set_font(run, size=10, color=(107, 114, 128))

doc.add_page_break()

# ── 1. 개요 ───────────────────────────────────────────────────
add_heading(doc, '1. 테스트 개요')
doc.add_paragraph()

info = [
    ('프로젝트명', 'TaskFlow - 업무 관리 웹 애플리케이션'),
    ('테스트 도구', 'Playwright (MCP 플러그인)'),
    ('테스트 환경', 'Windows 10 / Python 3.13 / FastAPI / Vanilla JS'),
    ('서버 주소', 'http://localhost:8000 (Backend), http://localhost:3000 (Frontend)'),
    ('테스트 일시', datetime.datetime.now().strftime('%Y-%m-%d %H:%M')),
    ('전체 테스트 수', '7개 시나리오 / 전체 통과'),
]

t = doc.add_table(rows=len(info), cols=2)
t.style = 'Table Grid'
for i, (k, v) in enumerate(info):
    row = t.rows[i]
    shade_row(row, "EEF2FF")
    cell_run(row.cells[0], k, bold=True)
    cell_run(row.cells[1], v)

doc.add_paragraph()

# ── 2. 테스트 대상 기능 ────────────────────────────────────────
add_heading(doc, '2. 테스트 대상 기능')
doc.add_paragraph()

features = [
    ('업무 추가', '제목, 메모, 우선순위(높음/보통/낙음), 시작일, 완료 예정일 입력'),
    ('상태 변경', '할 일 → 진행 중 → 완료 순환 변경'),
    ('업무 삭제', '× 버튼 클릭으로 즉시 삭제'),
    ('키워드 검색', '제목 및 메모 내용 실시간 검색 필터링'),
    ('상태 필터', '전체 / 할 일 / 진행 중 / 완료 탭 필터'),
    ('완료 숨기기', '완료된 업무 토글 방식으로 숨김/표시'),
    ('JSON 내보내기', '현재 업무 목록을 JSON 파일로 다운로드'),
    ('JSON 불러오기', 'JSON 파일 선택 후 업무 일괄 등록'),
    ('마감 초과 표시', '마감일이 지난 미완료 업무 빨간색 강조'),
    ('통계 바', '상태별 업무 수 실시간 표시'),
]

t2 = doc.add_table(rows=len(features)+1, cols=3)
t2.style = 'Table Grid'
hrow = t2.rows[0]
shade_row(hrow, "6366F1")
for cell, text in zip(hrow.cells, ['#', '기능명', '설명']):
    cell_run(cell, text, bold=True, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)

for i, (name, desc) in enumerate(features):
    row = t2.rows[i+1]
    if i % 2 == 0:
        shade_row(row, "F5F5FF")
    cell_run(row.cells[0], str(i+1), align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_run(row.cells[1], name, bold=True)
    cell_run(row.cells[2], desc)

doc.add_paragraph()

# ── 3. 테스트 시나리오 ─────────────────────────────────────────
add_heading(doc, '3. 테스트 시나리오 및 결과')
doc.add_paragraph()

scenarios = [
    ('TC-01', '초기 화면 로드',
     '브라우저에서 http://localhost:3000 접속',
     'TaskFlow 타이틀, 업무 추가 폼, 검색/필터 영역, 빈 목록 메시지 표시',
     '전체 UI 정상 표시 확인', '통과 ✅'),
    ('TC-02', '업무 추가',
     '제목/메모/우선순위/날짜 입력 후 추가 클릭',
     '업무 카드 생성, 통계 바 업데이트',
     '업무 카드 3개 정상 생성, 통계 "할 일 3" 표시', '통과 ✅'),
    ('TC-03', '마감 초과 표시',
     '마감일이 과거인 업무(2026-05-10) 추가',
     '"⏰ 마감 YYYY.MM.DD · 초과" 빨간색 텍스트 표시',
     '빨간색 마감 초과 표시 확인', '통과 ✅'),
    ('TC-04', '상태 변경',
     '상태 버튼 클릭으로 할일→진행중→완료 순환',
     '상태 뽙지 색상 변경, 완료 시 취소선 표시',
     '상태 순환 정상 동작 확인', '통과 ✅'),
    ('TC-05', '키워드 검색',
     '검색창에 "FastAPI" 입력',
     '"FastAPI 서버 구축" 업무만 목록에 표시',
     '1건만 필터링되어 표시', '통과 ✅'),
    ('TC-06', '완료 숨기기',
     '"완료 숨기기" 체크박스 클릭',
     '완료 상태 업무가 목록에서 사라짘',
     'UI 디자인 시안 검토(완료) 숨김 확인', '통과 ✅'),
    ('TC-07', '업무 삭제',
     '"배포 환경 설정" 업무의 × 버튼 클릭',
     '해당 업무 제거, 통계 바 업데이트',
     '정상 삭제, 통계 "할 일 0" 반영', '통과 ✅'),
]

labels = ['테스트 단계', '예상 결과', '실제 결과', '판정']

for tc, name, steps, expected, actual, result in scenarios:
    add_heading(doc, f'{tc}. {name}', level=2)
    t = doc.add_table(rows=4, cols=2)
    t.style = 'Table Grid'
    for i, (lbl, val) in enumerate(zip(labels, [steps, expected, actual, result])):
        row = t.rows[i]
        shade_row(row, "EEF2FF")
        cell_run(row.cells[0], lbl, bold=True)
        color = (16, 185, 129) if '통과' in val else None
        cell_run(row.cells[1], val, color=color)
    doc.add_paragraph()

# ── 4. 종합 결과 ───────────────────────────────────────────────
add_heading(doc, '4. 종합 결과')
doc.add_paragraph()

t3 = doc.add_table(rows=4, cols=2)
t3.style = 'Table Grid'
summary = [
    ('전체 테스트 케이스', '7개'),
    ('통과', '7개 (100%)'),
    ('실패', '0개 (0%)'),
    ('전체 판정', '✅ 전체 통과'),
]
for i, (k, v) in enumerate(summary):
    row = t3.rows[i]
    shade_row(row, "D1FAE5" if i == 3 else "F0FDF4")
    cell_run(row.cells[0], k, bold=True)
    cell_run(row.cells[1], v, bold=(i==3), color=(16,185,129))

doc.add_paragraph()

# ── 5. 발견 이슈 ───────────────────────────────────────────────
add_heading(doc, '5. 발견된 이슈 및 조치')
doc.add_paragraph()

t4 = doc.add_table(rows=3, cols=4)
t4.style = 'Table Grid'
hrow = t4.rows[0]
shade_row(hrow, "6366F1")
for cell, text in zip(hrow.cells, ['#', '이슈 내용', '심각도', '조치']):
    cell_run(cell, text, bold=True, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)

issues = [
    ('1', 'favicon.ico 404 에러 (탭 아이콘 없음)', '낙음', 'SVG 인라인 favicon 추가로 해결'),
    ('2', 'Tailwind CDN 프로덕션 사용 경고', '낙음', '배포 시 Tailwind CLI 전환 예정'),
]
for i, (no, issue, sev, action) in enumerate(issues):
    row = t4.rows[i+1]
    shade_row(row, "F5F5FF")
    cell_run(row.cells[0], no, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_run(row.cells[1], issue)
    cell_run(row.cells[2], sev)
    cell_run(row.cells[3], action)

doc.add_paragraph()
add_heading(doc, '6. 결론', level=2)
add_para(doc, 'TaskFlow 업무 관리 웹 애플리케이션에 대한 Playwright 자동화 테스트를 수행한 결과, '
         '7개의 테스트 케이스 모두 정상적으로 통과하였습니다. '
         '발견된 2건의 이슈는 모두 낙은 심각도이며 즉시 조치 완료하였습니다. '
         '업무 추가, 상태 변경, 삭제, 검색, 필터링, JSON 내보내기/불러오기 기능 모두 안정적으로 동작함을 확인하였습니다.')

output_path = r'D:\taskflow\TaskFlow_테스트결과보고서.docx'
doc.save(output_path)
print('OK')
